from docx import Document


def parse_docx(file_path: str) -> str:
    """Extract text from a .docx file including paragraphs and tables."""
    doc = Document(file_path)
    parts: list[str] = []

    for paragraph in doc.paragraphs:
        text = paragraph.text.strip()
        if text:
            parts.append(text)

    # Extract table content
    for table in doc.tables:
        for row in table.rows:
            row_text = " | ".join(cell.text.strip() for cell in row.cells)
            if row_text.strip(" |"):
                parts.append(row_text)

    return "\n\n".join(parts)
